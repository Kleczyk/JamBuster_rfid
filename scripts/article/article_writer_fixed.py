"""
Poprawiony moduł do tworzenia artykułów zgodnie z dokładnymi wymaganiami formatowania.
Implementuje wszystkie wymagania: marginesy, odstępy, dwie kolumny itp.
"""

from pathlib import Path

# Project root for default paths (when used as library from scripts/article/)
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, List, Dict, Any


class ArticleWriterFixed:
    """Klasa do tworzenia artykułów z dokładnym formatowaniem zgodnie z wymaganiami."""
    
    FONT_SIZES = {
        'title': 16,
        'authors': 10,
        'affiliation': 8,
        'doi': 8,
        'abstract': 7.5,
        'keywords': 8,
        'body': 9,
        'caption': 8,
        'table': 8,
        'literature': 8,
    }
    
    def __init__(self, docx_path: Optional[Path] = None):
        """Inicjalizacja ArticleWriterFixed."""
        if docx_path is None:
            docx_path = _PROJECT_ROOT / "article" / "Artykuł_poprawiony.docx"
        
        self.docx_path = Path(docx_path)
        self.doc = Document()
        self._two_column_section_created = False
        self._setup_default_formatting()
    
    def _setup_default_formatting(self):
        """Ustawienie formatowania strony zgodnie z wymaganiami."""
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)  # A4
        # Marginesy: górny, lewy, prawy 1,8 cm, dolny 2,5 cm
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.5)
    
    def _setup_two_columns(self, section, spacing_mm: float = 5.0):
        """Ustawienie dwóch kolumn w sekcji z odstępem 5mm."""
        sectPr = section._sectPr
        
        # Usuń istniejące kolumny jeśli są
        for existing_cols in sectPr.xpath('./w:cols'):
            sectPr.remove(existing_cols)
        
        # Utwórz element kolumn
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        # Odstęp 5mm w dwudziestkach punktu (twips)
        # 5mm = 0.5cm, 1cm = 28.35pt, więc 5mm = 14.175pt = 283.5 twips
        spacing_twips = int(Cm(spacing_mm / 10).pt * 20)
        cols.set(qn('w:space'), str(spacing_twips))
        cols.set(qn('w:equalWidth'), '1')
        
        # Dodaj kolumny do sekcji
        if not sectPr.xpath('./w:cols'):
            sectPr.append(cols)
        else:
            # Zastąp istniejące
            existing = sectPr.xpath('./w:cols')[0]
            sectPr.remove(existing)
            sectPr.append(cols)
    
    def add_authors(self, authors: List[str]) -> None:
        """Dodaj autorów (Arial bold 10, wyrównanie do prawej)."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        text = ", ".join(authors)
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['authors'])
        run.bold = True
    
    def add_affiliation(self, affiliation: str) -> None:
        """Dodaj afiliację (Arial 8, prawa, odstęp 1,5 linii)."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_before = Pt(18)  # 1.5 linii (12pt * 1.5)
        
        run = para.add_run(affiliation)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['affiliation'])
    
    def add_orcid(self, orcids: List[str]) -> None:
        """Dodaj ORCID (Arial 8, prawa)."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        text = "ORCID: " + "; ".join(orcids)
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['affiliation'])
    
    def add_doi(self, doi: str) -> None:
        """Dodaj DOI (Arial 8, lewa)."""
        para = self.doc.add_paragraph("doi: " + doi)
        run = para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['doi'])
    
    def add_title(self, title: str) -> None:
        """Dodaj tytuł (odstęp 8, Arial 16 bold, wyśrodkowany)."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Tytuł wyśrodkowany
        
        run = para.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['title'])
        run.bold = True
    
    def add_abstract(self, abstract_pl: str, abstract_en: str, title_en: str) -> None:
        """Dodaj streszczenie (dwa odstępy 8, Arial 7.5 kursywa, wyjustowane)."""
        # Streszczenie polskie - dwa odstępy 8
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(16)  # Dwa odstępy 8
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Wyjustowane
        
        run = para.add_run("Streszczenie. " + abstract_pl)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['abstract'])
        run.italic = True
        
        # Streszczenie angielskie - pojedynczy odstęp 8
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Wyjustowane
        
        run = para.add_run("Abstract. " + abstract_en + " (" + title_en + ")")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['abstract'])
        run.italic = True
    
    def add_keywords(self, keywords_pl: str, keywords_en: str) -> None:
        """Dodaj słowa kluczowe (jeden odstęp 8, Arial 8, wyjustowane)."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Wyjustowane
        
        text = f"Słowa kluczowe: {keywords_pl}\nKeywords: {keywords_en}."
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['keywords'])
    
    def _create_two_column_section(self):
        """Utwórz nową sekcję z dwoma kolumnami (odstęp 5mm)."""
        if self._two_column_section_created:
            return
        
        # Dodaj odstęp jednej interlinii (12pt)
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        
        # Utwórz nową sekcję z podziałem ciągłym
        # Używamy add_section() z odpowiednim typem
        new_section = self.doc.add_section()
        new_section.start_type = 3  # WD_SECTION.CONTINUOUS = 3
        
        # Ustaw dwie kolumny z odstępem 5mm
        self._setup_two_columns(new_section, spacing_mm=5.0)
        
        # Skopiuj marginesy z pierwszej sekcji
        first_section = self.doc.sections[0]
        new_section.left_margin = first_section.left_margin
        new_section.right_margin = first_section.right_margin
        new_section.top_margin = first_section.top_margin
        new_section.bottom_margin = first_section.bottom_margin
        
        self._two_column_section_created = True
    
    def add_heading(self, text: str, level: int = 2) -> None:
        """Dodaj nagłówek sekcji (Arial bold 9, bez numeracji)."""
        # Upewnij się, że sekcja dwukolumnowa jest utworzona
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_heading(text, level=level)
        
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
            run.bold = True
    
    def add_paragraph(self, text: str) -> None:
        """Dodaj akapit (Arial 9, wcięcie 5mm, wyjustowany)."""
        # Upewnij się, że sekcja dwukolumnowa jest utworzona
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Cm(0.5)  # 5mm
        
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
    
    def add_formula(self, formula_text: str, formula_number: int,
                   explanation: Optional[str] = None) -> None:
        """Dodaj wzór (odstęp 6pt góra/dół, numeracja po lewej, wzór na środku)."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Numeracja po lewej
        run = para.add_run(f"({formula_number})\t\t")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['body'])
        
        # Wzór (Times New Roman italic)
        run = para.add_run(formula_text)
        run.font.name = "Times New Roman"
        run.font.italic = True
        run.font.size = Pt(10)
        
        if explanation:
            para = self.doc.add_paragraph()
            run = para.add_run(explanation)
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
    
    def add_figure_caption(self, caption: str, figure_num: int) -> None:
        """Dodaj podpis rysunku (Arial 8, wyjustowany, bez kropki)."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        text = f"Rys.{figure_num}. {caption}"
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['caption'])
    
    def add_table(self, data: List[List[str]], title: str, table_num: int) -> None:
        """Dodaj tabelę z tytułem (Arial 8)."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        # Tytuł tabeli
        para = self.doc.add_paragraph()
        run = para.add_run(f"Tabela {table_num}. {title}")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['table'])
        
        if not data:
            return
        
        num_cols = len(data[0])
        table = self.doc.add_table(rows=len(data), cols=num_cols)
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_text)
                
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(self.FONT_SIZES['table'])
    
    def add_literature_item(self, item_num: int, citation: str) -> None:
        """Dodaj pozycję literatury (Arial 8, wysunięcie 5mm)."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(-0.5)  # Wysunięcie
        para.paragraph_format.left_indent = Cm(0.5)
        
        run = para.add_run(f"[{item_num}]\t{citation}")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['literature'])
    
    def add_author_info(self, authors_info: List[Dict[str, str]]) -> None:
        """Dodaj informacje o autorach przed literaturą."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        
        author_strings = []
        for author in authors_info:
            parts = [author.get('name', '')]
            if 'affiliation' in author:
                parts.append(author['affiliation'])
            if 'email' in author:
                parts.append(f"E-mail: {author['email']}")
            author_strings.append(", ".join(parts))
        
        text = "Autorzy: " + "; ".join(author_strings) + "."
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['body'])
    
    def add_literature_header(self) -> None:
        """Dodaj nagłówek LITERATURA."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph("LITERATURA")
        para.paragraph_format.space_before = Pt(12)
        
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
            run.bold = True
    
    def add_continuous_break_at_end(self) -> None:
        """Dodaj znak podziału ciągły na końcu (żeby kolumny miały jednakową długość)."""
        if not self._two_column_section_created:
            self._create_two_column_section()
        
        para = self.doc.add_paragraph()
        # Dodaj continuous break przez XML
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'continuous')
        para._p.append(br)
    
    def save(self, path: Optional[Path] = None) -> None:
        """Zapisz dokument."""
        save_path = path or self.docx_path
        self.doc.save(str(save_path))
        print(f"Dokument zapisany: {save_path}")


def open_article_fixed(docx_path: Optional[Path] = None) -> ArticleWriterFixed:
    """Otwórz artykuł do edycji."""
    return ArticleWriterFixed(docx_path)

