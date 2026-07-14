"""
Helper module for writing content to the article document (Artykuł.docx).

This module provides functions to add content to the article while maintaining
the document's formatting and style consistency according to the Polish academic
paper format described in the article itself.
"""

from pathlib import Path

# Project root for default paths (when used as library from scripts/article/)
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from typing import Optional, List, Dict, Any


class ArticleWriter:
    """Helper class for writing content to the article document following Polish academic format."""
    
    # Font sizes according to article specifications
    FONT_SIZES = {
        'title': 16,
        'authors': 10,
        'affiliation': 8,
        'doi': 8,
        'abstract': 7.5,
        'keywords': 8,
        'body': 9,
        'caption': 8,
        'table': 8,
        'literature': 8,
    }
    
    # Margins according to article specifications (in cm, converted to inches)
    MARGINS = {
        'top': 1.8,      # cm
        'left': 1.8,     # cm
        'right': 1.8,    # cm
        'bottom': 2.5,   # cm
    }
    
    def __init__(self, docx_path: Optional[Path] = None):
        """Initialize the ArticleWriter.
        
        Args:
            docx_path: Path to the article document. If None, uses default path.
        """
        if docx_path is None:
            docx_path = _PROJECT_ROOT / "article" / "Artykuł.docx"
        
        self.docx_path = Path(docx_path)
        
        if self.docx_path.exists():
            self.doc = Document(str(self.docx_path))
        else:
            # Create a new document if it doesn't exist
            self.doc = Document()
            self._setup_default_formatting()
    
    def _setup_default_formatting(self):
        """Set up default formatting according to article specifications."""
        # Set page margins (1.8 cm top/left/right, 2.5 cm bottom)
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69)  # A4 height
        section.left_margin = Cm(self.MARGINS['left'])
        section.right_margin = Cm(self.MARGINS['right'])
        section.top_margin = Cm(self.MARGINS['top'])
        section.bottom_margin = Cm(self.MARGINS['bottom'])
    
    def add_authors(self, authors: List[str], align_right: bool = True) -> None:
        """Add author names (Arial bold 10, right-aligned).
        
        Args:
            authors: List of author names with superscripts (e.g., "Daniel KLECZYŃSKI1")
            align_right: Whether to right-align the text
        """
        para = self.doc.add_paragraph()
        if align_right:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        text = ", ".join(authors)
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['authors'])
        run.bold = True
    
    def add_affiliation(self, affiliation: str, align_right: bool = True) -> None:
        """Add affiliation text (Arial 8, right-aligned).
        
        Args:
            affiliation: Affiliation text
            align_right: Whether to right-align the text
        """
        para = self.doc.add_paragraph()
        if align_right:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = para.add_run(affiliation)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['affiliation'])
    
    def add_orcid(self, orcids: List[str], align_right: bool = True) -> None:
        """Add ORCID information (Arial 8, right-aligned).
        
        Args:
            orcids: List of ORCID strings (e.g., ["1. 0009-0006-6814-7043", "2. ..."])
            align_right: Whether to right-align the text
        """
        para = self.doc.add_paragraph()
        if align_right:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        text = "ORCID: " + "; ".join(orcids)
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['affiliation'])
    
    def add_doi(self, doi: str) -> None:
        """Add DOI (Arial 8, left-aligned).
        
        Args:
            doi: DOI string
        """
        para = self.doc.add_paragraph("doi: " + doi)
        run = para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['doi'])
    
    def add_title(self, title: str, spacing_before: int = 8) -> None:
        """Add article title (Arial 16 bold).
        
        Args:
            title: Title text
            spacing_before: Spacing before title in points (default 8)
        """
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(spacing_before)
        
        run = para.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['title'])
        run.bold = True
    
    def add_abstract(self, abstract_pl: str, abstract_en: str, title_en: str) -> None:
        """Add abstract in Polish and English (Arial 7.5 italic).
        
        Args:
            abstract_pl: Polish abstract text
            abstract_en: English abstract text
            title_en: English title (will be added in parentheses)
        """
        # Polish abstract
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        
        run = para.add_run("Streszczenie. " + abstract_pl)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['abstract'])
        run.italic = True
        
        # English abstract
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        
        run = para.add_run("Abstract. " + abstract_en + " (" + title_en + ")")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['abstract'])
        run.italic = True
    
    def add_keywords(self, keywords_pl: str, keywords_en: str) -> None:
        """Add keywords in Polish and English (Arial 8).
        
        Args:
            keywords_pl: Polish keywords
            keywords_en: English keywords
        """
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        
        text = f"Słowa kluczowe: {keywords_pl}\nKeywords: {keywords_en}."
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['keywords'])
    
    def add_section_break(self) -> None:
        """Add continuous section break and switch to two-column layout."""
        para = self.doc.add_paragraph()
        para.add_run().add_break(WD_BREAK.CONTINUOUS)
        # Note: Column formatting needs to be done via section properties
        # This is a placeholder - actual column formatting requires section manipulation
    
    def add_heading(self, text: str, level: int = 2, 
                    bold: bool = True) -> None:
        """Add a section heading (Arial bold 9, no numbering).
        
        Args:
            text: Heading text
            level: Heading level (typically 2 for sections)
            bold: Whether to make text bold
        """
        para = self.doc.add_heading(text, level=level)
        
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
            run.bold = bold
    
    def add_paragraph(self, text: str, 
                     justify: bool = True,
                     first_line_indent: float = 0.5) -> None:
        """Add a body paragraph (Arial 9, first line indent 5mm, justified).
        
        Args:
            text: Paragraph text
            justify: Whether to justify the text
            first_line_indent: First line indent in cm (default 0.5cm = 5mm)
        """
        para = self.doc.add_paragraph(text)
        
        if justify:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set first line indent (5mm = 0.5cm)
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
        
        # Format runs
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
    
    def add_formula(self, formula_text: str, formula_number: int,
                   explanation: Optional[str] = None,
                   spacing: int = 6) -> None:
        """Add a formula with number and optional explanation.
        
        Args:
            formula_text: Formula text (will be inserted via equation editor)
            formula_number: Formula number for left-side numbering
            explanation: Optional explanation text (e.g., "gdzie: J – gęstość prądu...")
            spacing: Spacing above and below formula in points (default 6)
        """
        # Add spacing before
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(spacing)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Formula number (left-aligned)
        run = para.add_run(f"({formula_number})\t\t")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['body'])
        
        # Note: Actual formula insertion requires MathType or equation editor
        # This is a placeholder - formula_text should be inserted as equation object
        run = para.add_run(formula_text)
        run.font.name = "Times New Roman"
        run.font.italic = True
        run.font.size = Pt(10)
        
        # Add spacing after
        para.paragraph_format.space_after = Pt(spacing)
        
        # Add explanation if provided
        if explanation:
            para = self.doc.add_paragraph()
            run = para.add_run(explanation)
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
    
    def add_figure_caption(self, caption: str, figure_num: int) -> None:
        """Add a figure caption (Arial 8, justified, no period at end).
        
        Args:
            caption: Caption text (without "Rys." prefix)
            figure_num: Figure number
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        text = f"Rys.{figure_num}. {caption}"
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['caption'])
        # Note: No period at end as per article specifications
    
    def add_table(self, data: List[List[str]], 
                  title: str,
                  table_num: int) -> None:
        """Add a table with title (Arial 8).
        
        Args:
            data: List of rows, each row is a list of cell values
            title: Table title
            table_num: Table number
        """
        # Add table title
        para = self.doc.add_paragraph()
        run = para.add_run(f"Tabela {table_num}. {title}")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['table'])
        
        # Create table
        if not data:
            return
        
        num_cols = len(data[0])
        table = self.doc.add_table(rows=len(data), cols=num_cols)
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_text)
                
                # Format cell text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(self.FONT_SIZES['table'])
    
    def add_literature_item(self, item_num: int, citation: str) -> None:
        """Add a literature citation (Arial 8, hanging indent 5mm).
        
        Args:
            item_num: Citation number
            citation: Full citation text
        """
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(-0.5)  # Hanging indent
        para.paragraph_format.left_indent = Cm(0.5)
        
        run = para.add_run(f"[{item_num}]\t{citation}")
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['literature'])
    
    def add_author_info(self, authors_info: List[Dict[str, str]]) -> None:
        """Add author information at the end (before literature).
        
        Args:
            authors_info: List of dicts with keys: name, affiliation, email
        """
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        
        author_strings = []
        for author in authors_info:
            parts = [author.get('name', '')]
            if 'affiliation' in author:
                parts.append(author['affiliation'])
            if 'email' in author:
                parts.append(f"E-mail: {author['email']}")
            author_strings.append(", ".join(parts))
        
        text = "Autorzy: " + "; ".join(author_strings) + "."
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(self.FONT_SIZES['body'])
    
    def add_literature_header(self) -> None:
        """Add LITERATURA header."""
        para = self.doc.add_paragraph("LITERATURA")
        para.paragraph_format.space_before = Pt(12)
        
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(self.FONT_SIZES['body'])
            run.bold = True
    
    def add_page_break(self) -> None:
        """Add a page break."""
        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_break(6)  # Page break type
    
    def add_continuous_break(self) -> None:
        """Add a continuous section break."""
        # Add a paragraph with section break
        para = self.doc.add_paragraph()
        # Note: Section breaks are typically handled via section properties
        # This is a placeholder - actual section break insertion may require
        # direct XML manipulation or will be handled when setting up columns
    
    def save(self, path: Optional[Path] = None) -> None:
        """Save the document.
        
        Args:
            path: Path to save the document. If None, saves to original path.
        """
        save_path = path or self.docx_path
        self.doc.save(str(save_path))
        print(f"Document saved to: {save_path}")
    
    def get_document_info(self) -> Dict[str, Any]:
        """Get information about the document structure."""
        info = {
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'styles': [s.name for s in self.doc.styles],
            'sections': []
        }
        
        for section in self.doc.sections:
            info['sections'].append({
                'page_width': f"{section.page_width.inches:.2f}\"",
                'page_height': f"{section.page_height.inches:.2f}\"",
                'margins': {
                    'left': f"{section.left_margin.inches:.2f}\"",
                    'right': f"{section.right_margin.inches:.2f}\"",
                    'top': f"{section.top_margin.inches:.2f}\"",
                    'bottom': f"{section.bottom_margin.inches:.2f}\"",
                }
            })
        
        return info


# Convenience functions for quick access
def open_article(docx_path: Optional[Path] = None) -> ArticleWriter:
    """Open an article document for writing.
    
    Args:
        docx_path: Path to the article document. If None, uses default path.
    
    Returns:
        ArticleWriter instance
    """
    return ArticleWriter(docx_path)


if __name__ == "__main__":
    # Example usage
    writer = open_article()
    
    print("Article Writer - Polish Academic Format")
    print("=" * 50)
    print("\nDocument Info:")
    info = writer.get_document_info()
    print(f"  Paragraphs: {info['paragraphs']}")
    print(f"  Tables: {info['tables']}")
    print(f"  Available Styles: {', '.join(info['styles'][:10])}")
    
    print("\nUse this module to write content to the article:")
    print("  from article_writer import open_article")
    print("  writer = open_article()")
    print("  writer.add_heading('Nowa Sekcja', level=2)")
    print("  writer.add_paragraph('Tekst akapitu...')")
    print("  writer.save()")
