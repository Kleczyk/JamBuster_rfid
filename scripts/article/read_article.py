"""
Script to read and analyze the article document (Artykuł.docx).

This script extracts the content, structure, and formatting information
from the Word document to understand the article's format and principles.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def analyze_document(docx_path):
    """Analyze the structure and content of a Word document."""
    doc = Document(docx_path)
    
    print("=" * 80)
    print(f"ANALYZING DOCUMENT: {docx_path}")
    print("=" * 80)
    print()
    
    # Document properties
    print("DOCUMENT PROPERTIES:")
    print(f"  Core Properties:")
    print(f"    Title: {doc.core_properties.title or 'N/A'}")
    print(f"    Author: {doc.core_properties.author or 'N/A'}")
    print(f"    Subject: {doc.core_properties.subject or 'N/A'}")
    print(f"    Keywords: {doc.core_properties.keywords or 'N/A'}")
    print(f"    Created: {doc.core_properties.created or 'N/A'}")
    print(f"    Modified: {doc.core_properties.modified or 'N/A'}")
    print()
    
    # Section properties
    print("SECTION PROPERTIES:")
    for i, section in enumerate(doc.sections):
        print(f"  Section {i+1}:")
        print(f"    Page Width: {section.page_width.inches:.2f} inches")
        print(f"    Page Height: {section.page_height.inches:.2f} inches")
        print(f"    Left Margin: {section.left_margin.inches:.2f} inches")
        print(f"    Right Margin: {section.right_margin.inches:.2f} inches")
        print(f"    Top Margin: {section.top_margin.inches:.2f} inches")
        print(f"    Bottom Margin: {section.bottom_margin.inches:.2f} inches")
        print(f"    Header Distance: {section.header_distance.inches:.2f} inches")
        print(f"    Footer Distance: {section.footer_distance.inches:.2f} inches")
    print()
    
    # Paragraph analysis
    print("PARAGRAPH STRUCTURE:")
    print("-" * 80)
    
    paragraph_styles = {}
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Only analyze non-empty paragraphs
            style_name = para.style.name
            if style_name not in paragraph_styles:
                paragraph_styles[style_name] = []
            
            # Extract formatting information
            para_info = {
                'index': i,
                'text': para.text[:100] + '...' if len(para.text) > 100 else para.text,
                'alignment': str(para.alignment),
                'runs': []
            }
            
            # Analyze runs (text with specific formatting)
            for run in para.runs:
                run_info = {
                    'text': run.text[:50] + '...' if len(run.text) > 50 else run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None,
                }
                para_info['runs'].append(run_info)
            
            paragraph_styles[style_name].append(para_info)
    
    # Print style summary
    for style_name, paragraphs in paragraph_styles.items():
        print(f"\n  Style: {style_name} ({len(paragraphs)} paragraphs)")
        for para_info in paragraphs[:3]:  # Show first 3 examples
            print(f"    [{para_info['index']}] Alignment: {para_info['alignment']}")
            print(f"        Text: {para_info['text']}")
            if para_info['runs']:
                print(f"        Runs: {len(para_info['runs'])}")
                for run_info in para_info['runs'][:2]:  # Show first 2 runs
                    fmt = []
                    if run_info['bold']: fmt.append('Bold')
                    if run_info['italic']: fmt.append('Italic')
                    if run_info['underline']: fmt.append('Underline')
                    fmt_str = ', '.join(fmt) if fmt else 'Normal'
                    font_info = f"Font: {run_info['font_name']}"
                    if run_info['font_size']:
                        font_info += f", Size: {run_info['font_size']}pt"
                    print(f"          - {run_info['text'][:40]} ({fmt_str}, {font_info})")
        if len(paragraphs) > 3:
            print(f"        ... and {len(paragraphs) - 3} more paragraphs")
    
    print()
    
    # Table analysis
    print("TABLES:")
    print("-" * 80)
    if doc.tables:
        for i, table in enumerate(doc.tables):
            print(f"\n  Table {i+1}:")
            print(f"    Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            # Show first few cells
            for row_idx, row in enumerate(table.rows[:3]):
                cells_text = [cell.text[:30] for cell in row.cells]
                print(f"    Row {row_idx+1}: {' | '.join(cells_text)}")
            if len(table.rows) > 3:
                print(f"    ... and {len(table.rows) - 3} more rows")
    else:
        print("  No tables found")
    print()
    
    # Full text extraction
    print("FULL TEXT CONTENT:")
    print("-" * 80)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    print("\n".join(full_text))
    print()
    
    return {
        'document': doc,
        'styles': paragraph_styles,
        'full_text': full_text,
        'tables': doc.tables
    }


def extract_formatting_template(docx_path):
    """Extract formatting information to create a template for writing."""
    doc = Document(docx_path)
    
    template = {
        'sections': [],
        'styles': {},
        'default_fonts': {}
    }
    
    # Section formatting
    for section in doc.sections:
        template['sections'].append({
            'page_width': section.page_width,
            'page_height': section.page_height,
            'margins': {
                'left': section.left_margin,
                'right': section.right_margin,
                'top': section.top_margin,
                'bottom': section.bottom_margin,
            }
        })
    
    # Style information
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name
            if style_name not in template['styles']:
                template['styles'][style_name] = {
                    'alignment': para.alignment,
                    'runs': []
                }
                
                # Get run formatting
                for run in para.runs:
                    if run.text.strip():
                        run_fmt = {
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                        }
                        if run_fmt not in template['styles'][style_name]['runs']:
                            template['styles'][style_name]['runs'].append(run_fmt)
                        break  # Just get first run as template
    
    return template


def print_formatting_guide(template):
    """Print a guide on how to format new content."""
    print("\n" + "=" * 80)
    print("FORMATTING TEMPLATE GUIDE")
    print("=" * 80)
    print("\nUse this information to maintain consistent formatting when writing:")
    print()
    
    if template['sections']:
        section = template['sections'][0]
        print("Page Setup:")
        print(f"  Page Size: {section['page_width'].inches:.2f}\" x {section['page_height'].inches:.2f}\"")
        print(f"  Margins: L={section['margins']['left'].inches:.2f}\", "
              f"R={section['margins']['right'].inches:.2f}\", "
              f"T={section['margins']['top'].inches:.2f}\", "
              f"B={section['margins']['bottom'].inches:.2f}\"")
        print()
    
    if template['styles']:
        print("Available Styles:")
        for style_name, style_info in template['styles'].items():
            print(f"  {style_name}:")
            print(f"    Alignment: {style_info['alignment']}")
            if style_info['runs']:
                for run_fmt in style_info['runs']:
                    print(f"    Font: {run_fmt['font_name']}")
                    if run_fmt['font_size']:
                        print(f"    Size: {run_fmt['font_size'].pt}pt")
                    fmt_parts = []
                    if run_fmt['bold']: fmt_parts.append('Bold')
                    if run_fmt['italic']: fmt_parts.append('Italic')
                    if run_fmt['underline']: fmt_parts.append('Underline')
                    if fmt_parts:
                        print(f"    Formatting: {', '.join(fmt_parts)}")
        print()


if __name__ == "__main__":
    from _paths import PROJECT_ROOT
    # Path to the article document
    article_path = PROJECT_ROOT / "article" / "Artykuł.docx"
    
    if not article_path.exists():
        print(f"ERROR: Article document not found at {article_path}")
        print("Please ensure the file exists.")
        exit(1)
    
    # Analyze the document
    analysis = analyze_document(article_path)
    
    # Extract formatting template
    template = extract_formatting_template(article_path)
    
    # Print formatting guide
    print_formatting_guide(template)
    
    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE")
    print("=" * 80)
    print("\nYou can now use this information to:")
    print("  1. Understand the article's structure and formatting")
    print("  2. Maintain consistent formatting when adding new content")
    print("  3. Create helper functions to write content in the same style")


















